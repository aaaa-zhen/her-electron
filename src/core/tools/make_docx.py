#!/usr/bin/env python3
"""Generate a polished DOCX from JSON input.

Usage: echo '<json>' | python3 make_docx.py <output_path>

JSON schema:
{
  "title": "Document Title",
  "subtitle": "Optional subtitle",
  "author": "Author name",
  "theme": "default" | "formal" | "modern" | "minimal",
  "header_text": "Optional header on every page",
  "footer_text": "Optional footer on every page",
  "sections": [
    {
      "type": "heading" | "paragraph" | "bullet_list" | "numbered_list" | "table" | "quote" | "code" | "page_break" | "image",
      "level": 1,
      "text": "Section text",
      "items": ["item1", "item2"],
      "rows": [["a","b"],["c","d"]],
      "headers": ["Col A", "Col B"],
      "language": "python",
      "src": "/path/to/image.png",
      "width_inches": 5
    }
  ]
}
"""

import json
import sys
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

THEMES = {
    "default": {
        "title_color": RGBColor(0x1a, 0x1a, 0x2e),
        "heading_color": RGBColor(0x1a, 0x1a, 0x2e),
        "body_color": RGBColor(0x33, 0x33, 0x33),
        "accent": RGBColor(0x10, 0xb9, 0x81),
        "quote_bg": "F0F7F4",
        "quote_border": "10B981",
        "code_bg": "F5F5F5",
        "table_header_bg": "10B981",
        "table_header_fg": RGBColor(0xff, 0xff, 0xff),
        "table_alt_bg": "F0F7F4",
        "font": "Helvetica Neue",
        "cjk_font": "PingFang SC",
    },
    "formal": {
        "title_color": RGBColor(0x1a, 0x1a, 0x1a),
        "heading_color": RGBColor(0x1a, 0x1a, 0x1a),
        "body_color": RGBColor(0x33, 0x33, 0x33),
        "accent": RGBColor(0x2c, 0x3e, 0x50),
        "quote_bg": "F4F4F4",
        "quote_border": "2C3E50",
        "code_bg": "F8F8F8",
        "table_header_bg": "2C3E50",
        "table_header_fg": RGBColor(0xff, 0xff, 0xff),
        "table_alt_bg": "F4F6F7",
        "font": "Times New Roman",
        "cjk_font": "Songti SC",
    },
    "modern": {
        "title_color": RGBColor(0x0f, 0x17, 0x2a),
        "heading_color": RGBColor(0x0f, 0x17, 0x2a),
        "body_color": RGBColor(0x37, 0x41, 0x51),
        "accent": RGBColor(0x60, 0xa5, 0xfa),
        "quote_bg": "EFF6FF",
        "quote_border": "60A5FA",
        "code_bg": "F1F5F9",
        "table_header_bg": "1E293B",
        "table_header_fg": RGBColor(0xff, 0xff, 0xff),
        "table_alt_bg": "F1F5F9",
        "font": "Helvetica Neue",
        "cjk_font": "PingFang SC",
    },
    "minimal": {
        "title_color": RGBColor(0x11, 0x11, 0x11),
        "heading_color": RGBColor(0x11, 0x11, 0x11),
        "body_color": RGBColor(0x44, 0x44, 0x44),
        "accent": RGBColor(0x55, 0x55, 0x55),
        "quote_bg": "FAFAFA",
        "quote_border": "CCCCCC",
        "code_bg": "FAFAFA",
        "table_header_bg": "333333",
        "table_header_fg": RGBColor(0xff, 0xff, 0xff),
        "table_alt_bg": "FAFAFA",
        "font": "Helvetica Neue",
        "cjk_font": "PingFang SC",
    },
}


def set_run_font(run, font_name, cjk_font, size, color=None, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), cjk_font)


def add_heading(doc, text, level, theme):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, theme["font"], theme["cjk_font"],
                     {1: 22, 2: 16, 3: 13}.get(level, 12),
                     color=theme["heading_color"], bold=True)
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_paragraph(doc, text, theme, alignment=None, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, theme["font"], theme["cjk_font"], 11, color=theme["body_color"],
                 bold=bold, italic=italic)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    if alignment == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return p


def add_bullet_list(doc, items, theme):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(str(item))
        set_run_font(run, theme["font"], theme["cjk_font"], 11, color=theme["body_color"])
        p.paragraph_format.space_after = Pt(3)


def add_numbered_list(doc, items, theme):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(str(item))
        set_run_font(run, theme["font"], theme["cjk_font"], 11, color=theme["body_color"])
        p.paragraph_format.space_after = Pt(3)


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, theme):
    col_count = len(headers) if headers else (len(rows[0]) if rows else 0)
    if col_count == 0:
        return
    row_count = (1 if headers else 0) + len(rows)
    table = doc.add_table(rows=row_count, cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    start_row = 0
    if headers:
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(h))
            set_run_font(run, theme["font"], theme["cjk_font"], 10,
                         color=theme["table_header_fg"], bold=True)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, theme["table_header_bg"])
        start_row = 1

    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            if ci >= col_count:
                break
            cell = table.rows[start_row + ri].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, theme["font"], theme["cjk_font"], 10, color=theme["body_color"])
            if ri % 2 == 1:
                set_cell_shading(cell, theme["table_alt_bg"])

    doc.add_paragraph()  # spacing after table


def add_quote(doc, text, theme):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(1.5)
    pf.space_before = Pt(8)
    pf.space_after = Pt(8)
    pf.line_spacing = Pt(20)

    # Left border via XML
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="12" w:space="8" w:color="{theme["quote_border"]}"/>'
        f'</w:pBdr>'
    )
    p._element.get_or_add_pPr().append(pBdr)

    # Background shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{theme["quote_bg"]}" w:val="clear"/>')
    p._element.get_or_add_pPr().append(shading)

    run = p.add_run(text)
    set_run_font(run, theme["font"], theme["cjk_font"], 11,
                 color=theme["body_color"], italic=True)


def add_code(doc, text, theme, language=""):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.5)
    pf.right_indent = Cm(0.5)
    pf.space_before = Pt(8)
    pf.space_after = Pt(8)

    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{theme["code_bg"]}" w:val="clear"/>')
    p._element.get_or_add_pPr().append(shading)

    run = p.add_run(text)
    run.font.name = "Menlo"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    r = run._element
    r.rPr.rFonts.set(qn("w:eastAsia"), "Menlo")


def add_image(doc, src, width_inches=5):
    if not os.path.isfile(src):
        doc.add_paragraph(f"[Image not found: {src}]")
        return
    doc.add_picture(src, width=Inches(width_inches))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def generate(data, output_path):
    doc = Document()

    theme_name = data.get("theme", "default")
    theme = THEMES.get(theme_name, THEMES["default"])

    # Default styles
    style = doc.styles["Normal"]
    style.font.name = theme["font"]
    style.font.size = Pt(11)
    style.font.color.rgb = theme["body_color"]
    style.element.rPr.rFonts.set(qn("w:eastAsia"), theme["cjk_font"])

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Header
    header_text = data.get("header_text", "")
    if header_text:
        for section in doc.sections:
            header = section.header
            header.is_linked_to_previous = False
            hp = header.paragraphs[0]
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            hr = hp.add_run(header_text)
            set_run_font(hr, theme["font"], theme["cjk_font"], 8,
                         color=RGBColor(0x99, 0x99, 0x99))

    # Footer
    footer_text = data.get("footer_text", "")
    if footer_text:
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            fp = footer.paragraphs[0]
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fr = fp.add_run(footer_text)
            set_run_font(fr, theme["font"], theme["cjk_font"], 8,
                         color=RGBColor(0x99, 0x99, 0x99))

    # Title
    title = data.get("title", "")
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        set_run_font(run, theme["font"], theme["cjk_font"], 26,
                     color=theme["title_color"], bold=True)
        p.paragraph_format.space_after = Pt(4)

    # Subtitle
    subtitle = data.get("subtitle", "")
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(subtitle)
        set_run_font(run, theme["font"], theme["cjk_font"], 14,
                     color=RGBColor(0x88, 0x88, 0x88), italic=True)
        p.paragraph_format.space_after = Pt(20)

    # Sections
    for sec in data.get("sections", []):
        sec_type = sec.get("type", "paragraph")

        if sec_type == "heading":
            level = sec.get("level", 1)
            add_heading(doc, sec.get("text", ""), min(level, 3), theme)

        elif sec_type == "paragraph":
            text = sec.get("text", "")
            for line in text.split("\n"):
                if line.strip():
                    add_paragraph(doc, line, theme,
                                  alignment=sec.get("alignment"),
                                  bold=sec.get("bold", False),
                                  italic=sec.get("italic", False))

        elif sec_type == "bullet_list":
            items = sec.get("items", [])
            add_bullet_list(doc, items, theme)

        elif sec_type == "numbered_list":
            items = sec.get("items", [])
            add_numbered_list(doc, items, theme)

        elif sec_type == "table":
            headers = sec.get("headers", [])
            rows = sec.get("rows", [])
            add_table(doc, headers, rows, theme)

        elif sec_type == "quote":
            add_quote(doc, sec.get("text", ""), theme)

        elif sec_type == "code":
            add_code(doc, sec.get("text", ""), theme, language=sec.get("language", ""))

        elif sec_type == "page_break":
            doc.add_page_break()

        elif sec_type == "image":
            add_image(doc, sec.get("src", ""), sec.get("width_inches", 5))

    doc.save(output_path)
    print(f"OK:{output_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: echo '<json>' | python3 make_docx.py <output>", file=sys.stderr)
        sys.exit(1)
    raw = sys.stdin.read()
    data = json.loads(raw)
    generate(data, sys.argv[1])
